import docx

d = docx.Document('demo.docx')
print(d.paragraphs)
print()
print(d.paragraphs[0].text)
print()
p = d.paragraphs[1]
print(p.runs)
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
p.runs[3].underline = True
p.runs[3].text = "this is underlined and italic"
d.save("demo2.docx")

d = docx.Document()
d.add_paragraph("Hello world, I am Prajwal!")
d.save("demo4.docx")
p = d.paragraphs[0]
p.add_run('New run.')
p.runs[1].bold = True
d.save("demo4.docx")
